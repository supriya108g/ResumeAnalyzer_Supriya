import json
import os
import re
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Dict
from xml.etree import ElementTree

import requests
import streamlit as st

from ResumeAnalyzer import extract_candidate_name, is_likely_candidate_name

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None
    letter = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

FASTAPI_URL = os.getenv("FASTAPI_URL", "http://127.0.0.1:8001/resume-analyzer")
LATEST_ANALYSIS_FILE = Path(__file__).resolve().parent / "latest_analysis.json"


def extract_candidate_name_from_resume(resume_text: str) -> str:
    return extract_candidate_name(resume_text)


def read_docx_text(data: bytes) -> str:
    # Read WordprocessingML directly because python-docx omits text-box content.
    word_namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraph_tag = f"{word_namespace}p"
    text_tag = f"{word_namespace}t"
    lines = []

    with zipfile.ZipFile(BytesIO(data)) as archive:
        part_names = ["word/document.xml"]
        part_names.extend(
            name
            for name in archive.namelist()
            if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
        )
        for part_name in part_names:
            if part_name not in archive.namelist():
                continue
            root = ElementTree.fromstring(archive.read(part_name))
            for paragraph in root.iter(paragraph_tag):
                # A text box may nest paragraphs; process only leaf nodes to avoid duplicates.
                nested_paragraphs = list(paragraph.iter(paragraph_tag))[1:]
                if nested_paragraphs:
                    continue
                line = "".join(
                    node.text or ""
                    for node in paragraph.iter(text_tag)
                ).strip()
                if line and (not lines or line != lines[-1]):
                    lines.append(line)

    return "\n".join(lines)


def read_text_from_file(file_obj) -> str:
    filename = file_obj.name.lower()

    if filename.endswith(".txt"):
        return file_obj.read().decode("utf-8", errors="replace")

    if filename.endswith(".pdf"):
        if PdfReader is None:
            raise RuntimeError("pypdf is not installed. Please install it from requirements.txt")
        pdf_reader = PdfReader(BytesIO(file_obj.read()))
        pages = []
        for page in pdf_reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages)

    if filename.endswith(".docx"):
        return read_docx_text(file_obj.getvalue())

    raise ValueError("Unsupported file type. Please upload a .txt, .pdf, or .docx file.")


def call_resume_api(candidate_name: str, resume_text: str, job_description: str) -> Dict:
    payload = {
        "candidate_name": candidate_name,
        "resume_text": resume_text,
        "job_description": job_description,
        "user_id": "streamlit-user",
    }

    response = requests.post(FASTAPI_URL, json=payload, timeout=120)
    response.raise_for_status()
    return response.json()


def clean_text(value: str) -> str:
    return (value or "").strip()


def save_latest_analysis(result: Dict) -> None:
    LATEST_ANALYSIS_FILE.write_text(json.dumps(result, indent=2), encoding="utf-8")


def load_latest_analysis() -> Dict:
    if not LATEST_ANALYSIS_FILE.exists():
        return {}
    try:
        return json.loads(LATEST_ANALYSIS_FILE.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}


def build_resume_export_text(generated_resume: Dict) -> str:
    if not generated_resume:
        return "Improved resume content is unavailable."

    sections = [
        "PROFESSIONAL SUMMARY",
        str(generated_resume.get("professional_summary", "")).strip(),
        "",
        "SKILLS",
        str(generated_resume.get("skills_section", "")).strip(),
        "",
        "EXPERIENCE",
    ]

    for bullet in generated_resume.get("experience_bullets", []):
        sections.append(f"- {bullet}")

    sections.extend(["", "PROJECTS"])
    for bullet in generated_resume.get("project_descriptions", []):
        sections.append(f"- {bullet}")

    sections.extend(["", "EDUCATION", str(generated_resume.get("education_section", "")).strip()])
    return "\n".join(sections)


def build_resume_docx_bytes(generated_resume: Dict) -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install it from requirements.txt")

    document = Document()
    document.add_heading(str(generated_resume.get("candidate_name", "Candidate")), level=1)

    summary = generated_resume.get("professional_summary", "")
    if summary:
        document.add_paragraph(summary)

    skills = generated_resume.get("skills_section", "")
    if skills:
        document.add_heading("Skills", level=2)
        document.add_paragraph(str(skills))

    experience = generated_resume.get("experience_bullets", [])
    if experience:
        document.add_heading("Experience", level=2)
        for bullet in experience:
            document.add_paragraph(bullet, style="List Bullet")

    projects = generated_resume.get("project_descriptions", [])
    if projects:
        document.add_heading("Projects", level=2)
        for bullet in projects:
            document.add_paragraph(bullet, style="List Bullet")

    education = generated_resume.get("education_section", "")
    if education:
        document.add_heading("Education", level=2)
        document.add_paragraph(education)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_resume_pdf_bytes(generated_resume: Dict) -> bytes:
    if canvas is None or letter is None:
        raise RuntimeError("reportlab is not installed. Please install it from requirements.txt")

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    candidate_name = str(generated_resume.get("candidate_name", "Candidate"))
    pdf.setTitle(f"{candidate_name} Resume")
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(72, 760, candidate_name)

    y = 720
    lines = build_resume_export_text(generated_resume).splitlines()
    pdf.setFont("Helvetica", 10)

    for line in lines:
        if y < 60:
            pdf.showPage()
            y = 760
            pdf.setFont("Helvetica", 10)
        pdf.drawString(72, y, line[:100])
        y -= 18

    pdf.save()
    return buffer.getvalue()


def render_analysis_result(result: Dict) -> None:
    if result.get("reason") == "missing_openai_api_key":
        st.error(result.get("message", "OPENAI_API_KEY must be configured to continue."))
        return

    st.subheader("Result Summary")
    if result.get("used_existing_log"):
        st.info(result.get("message", "Improved Resume already exists for the job description"))
        st.write(f"Decision: {result.get('decision', 'Reuse existing improved resume')}")
        st.write(f"Status: {result.get('status', 'Completed')}")
    else:
        col_summary_1, col_summary_2, col_summary_3, col_summary_4 = st.columns(4)
        with col_summary_1:
            st.metric("Uploaded Resume Score", result.get("uploaded_resume_score_out_of_10", "N/A"))
        with col_summary_2:
            st.metric("Improved Resume Score", result.get("improved_resume_score_out_of_10", "N/A"))
        with col_summary_3:
            st.markdown(f"**Decision**\n\n{result.get('decision', 'N/A')}")
        with col_summary_4:
            st.metric("Status", result.get("status", "N/A"))

    if result.get("message") and not result.get("used_existing_log"):
        st.info(result["message"])

    st.subheader("Generated Resume")
    generated = result.get("generated_resume", {})
    if not generated:
        return

    candidate_name = clean_text(generated.get("candidate_name", result.get("candidate_name", "")))
    if not is_likely_candidate_name(candidate_name):
        candidate_name = ""
    if candidate_name:
        st.title(candidate_name)

    st.write("### Professional Summary")
    st.write(generated.get("professional_summary", ""))

    st.write("### Skills")
    skills_value = generated.get("skills_section", "")
    st.write(", ".join(skills_value) if isinstance(skills_value, list) else skills_value)

    st.write("### Experience")
    for bullet in generated.get("experience_bullets", []):
        st.write("- " + bullet)

    st.write("### Projects")
    for bullet in generated.get("project_descriptions", []):
        st.write("- " + bullet)

    education = clean_text(generated.get("education_section", ""))
    st.write("### Education")
    st.write(education or "Education not provided in uploaded resume.")

    st.subheader("Download Improved Resume")
    filename_base = candidate_name.replace(" ", "_") if candidate_name else "improved_resume"
    st.download_button(
        label="Download as DOCX",
        data=build_resume_docx_bytes(generated),
        file_name=f"{filename_base}_improved_resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        on_click="ignore",
    )
    st.download_button(
        label="Download as PDF",
        data=build_resume_pdf_bytes(generated),
        file_name=f"{filename_base}_improved_resume.pdf",
        mime="application/pdf",
        on_click="ignore",
    )


st.set_page_config(page_title="Enterprise AI Resume Generator", page_icon="📄", layout="wide")

if st.query_params.get("view") == "agent-analysis":
    if st.sidebar.button("Resume Analyzer"):
        st.query_params.clear()
        st.rerun()
    st.title("Agent Analysis")
    analysis_result = st.session_state.get("analysis_result") or load_latest_analysis()
    if not analysis_result:
        st.info("Run a resume analysis first to view its details.")
    else:
        agent_outputs = analysis_result.get("crewai", {}).get("agent_outputs", {})
        if agent_outputs:
            labels = {
                "profile_analyzer_agent": "Profile Analyzer Agent",
                "ats_optimization_agent": "ATS Optimization Agent",
                "resume_writer_agent": "Resume Writer Agent",
                "reviewer_agent": "Reviewer Agent",
            }
            for key, label in labels.items():
                st.subheader(label)
                st.json(agent_outputs.get(key, {}))
        else:
            st.info("Individual CrewAI outputs are unavailable for this older analysis. Run the analysis again.")
    st.stop()

st.title("Enterprise AI Resume Generator")
st.caption("Upload a resume and a job description, then analyze and improve ATS alignment using FastAPI + CrewAI + RAG workflow.")

with st.form("resume_analysis_form"):
    col1, col2 = st.columns(2)
    with col1:
        resume_file = st.file_uploader("Upload Resume", type=["txt", "pdf", "docx"])

    with col2:
        jd_file = st.file_uploader("Upload Job Description", type=["txt", "pdf", "docx"])

    submitted = st.form_submit_button("Analyze Resume", type="primary")

input_fingerprint = (
    (resume_file.name, resume_file.size) if resume_file is not None else None,
    (jd_file.name, jd_file.size) if jd_file is not None else None,
)
# Results created before fingerprint tracking cannot be tied to the current files.
if "analysis_result" in st.session_state and "analyzed_input_fingerprint" not in st.session_state:
    st.session_state.pop("analysis_result", None)
# Selecting different inputs clears the display but preserves results during navigation.
if any(input_fingerprint) and input_fingerprint != st.session_state.get("analyzed_input_fingerprint"):
    st.session_state.pop("analysis_result", None)

if submitted:
    try:
        resume_text = ""
        job_description = ""
        if resume_file is not None:
            resume_text = read_text_from_file(resume_file)
        if jd_file is not None:
            job_description = read_text_from_file(jd_file)

        resume_text = clean_text(resume_text)
        job_description = clean_text(job_description)

        if not resume_text:
            st.warning("Please provide resume content before analysis.")
        elif not job_description:
            st.warning("Please provide a job description before analysis.")
        else:
            with st.spinner("Running resume analysis and ATS optimization..."):
                candidate_name = extract_candidate_name_from_resume(resume_text)
                result = call_resume_api(candidate_name, resume_text, job_description)
                st.session_state["analysis_result"] = result
                st.session_state["analyzed_input_fingerprint"] = input_fingerprint
                save_latest_analysis(result)

            if result.get("reason") != "missing_openai_api_key":
                st.success("Analysis completed successfully.")

    except requests.exceptions.RequestException as exc:
        st.error(f"FastAPI connection failed. Please start the server first. Error: {exc}")
    except Exception as exc:
        st.error(f"Processing failed: {exc}")

analysis_result = st.session_state.get("analysis_result")
if analysis_result:
    render_analysis_result(analysis_result)

if st.sidebar.button("Agent Analysis"):
    st.query_params["view"] = "agent-analysis"
    st.rerun()
